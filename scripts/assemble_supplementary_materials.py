#!/usr/bin/env python3
"""Assemble the final supplementary-figure DOCX from the immutable template.

The source DOCX is never written. Existing figure blocks are cloned from it,
revision figures are inserted using its image and caption formatting, and the
93 consensus-locus GRCh38-versus-panTro6 SVbyEye plots are appended as a
two-plots-per-page portrait appendix rather than promoted to numbered
supplementary figures.
"""

from __future__ import annotations

import argparse
import copy
import csv
import hashlib
import io
import json
import os
import re
import sys
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pymupdf
from PIL import Image
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "stats"))

from supplementary_inventory import (  # noqa: E402
    FINAL_SUPPLEMENTARY_FIGURES,
    SUPPLEMENT_TEMPLATE,
    SUPPLEMENT_TEMPLATE_SHA256,
    SVBYEYE_APPENDIX_TITLE,
    SVBYEYE_CONSENSUS_PDF,
)


FIGURE_RE = re.compile(r"^Figure S(\d+)\.")
APPENDIX_CAPTION_RE = re.compile(r"^SVbyEye alignment (\d+) of 93\.")
CALL_LABELS = {
    "direct": "ancestral",
    "inverted": "derived",
    "na": "not callable",
}
FIGURE_FRAME_WIDTH_IN = 6.5
FIGURE_FRAME_HEIGHT_IN = 5.75
FIGURE_FRAME_WIDTH_PX = 2400
FIGURE_FRAME_HEIGHT_PX = round(
    FIGURE_FRAME_WIDTH_PX * FIGURE_FRAME_HEIGHT_IN / FIGURE_FRAME_WIDTH_IN
)
SVBYEYE_PLOTS_PER_PAGE = 2
SVBYEYE_FRAME_WIDTH_IN = FIGURE_FRAME_WIDTH_IN

IMMUTABLE_TEMPLATE_PARTS = (
    "word/styles.xml",
    "word/settings.xml",
    "word/fontTable.xml",
    "word/theme/theme1.xml",
    "word/header1.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def restore_template_parts(document_path: Path, template: Path) -> None:
    """Put immutable style/header parts back byte-for-byte after python-docx saves."""
    rewritten = document_path.with_suffix(".restored.docx")
    with ZipFile(document_path) as generated, ZipFile(template) as source, ZipFile(
        rewritten, "w", compression=ZIP_DEFLATED
    ) as output:
        generated_names = set(generated.namelist())
        source_names = set(source.namelist())
        missing = set(IMMUTABLE_TEMPLATE_PARTS) - generated_names
        if missing:
            raise RuntimeError(f"Generated DOCX is missing template parts: {sorted(missing)}")
        for info in generated.infolist():
            name = info.filename
            if name in IMMUTABLE_TEMPLATE_PARTS and name in source_names:
                payload = source.read(name)
            elif re.fullmatch(r"word/media/image(?:[1-9]|1[0-3])\.png", name):
                payload = source.read(name)
            else:
                payload = generated.read(name)
            output.writestr(info, payload)
    os.replace(rewritten, document_path)


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.xpath(".//w:t"))


def contains_drawing(element) -> bool:
    return bool(element.xpath(".//w:drawing"))


def replace_prefix(element, old: str, new: str) -> None:
    nodes = element.xpath(".//w:t")
    full = "".join(node.text or "" for node in nodes)
    if not full.startswith(old):
        raise RuntimeError(f"Expected paragraph prefix {old!r}, observed {full[:80]!r}")

    remaining = len(old)
    inserted = False
    for node in nodes:
        text = node.text or ""
        if remaining <= 0:
            break
        take = min(len(text), remaining)
        tail = text[take:]
        if not inserted:
            node.text = new + tail
            inserted = True
        else:
            node.text = tail
        remaining -= take
    if remaining:
        raise RuntimeError(f"Could not replace complete prefix {old!r}")


def replace_paragraph_text(paragraph, old: str, new: str) -> None:
    if paragraph.text.strip() != old:
        raise RuntimeError(
            f"Expected front-matter paragraph {old!r}, observed {paragraph.text!r}"
        )
    replace_prefix(paragraph._p, old, new)


def copy_paragraph_properties(source, target) -> None:
    source_ppr = source._p.pPr
    target_ppr = target._p.pPr
    if target_ppr is not None:
        target._p.remove(target_ppr)
    if source_ppr is not None:
        target._p.insert(0, copy.deepcopy(source_ppr))


def copy_run_properties(source_run, target_run) -> None:
    source_rpr = source_run._r.rPr
    target_rpr = target_run._r.rPr
    if target_rpr is not None:
        target_run._r.remove(target_rpr)
    if source_rpr is not None:
        target_run._r.insert(0, copy.deepcopy(source_rpr))


def template_runs(caption_template):
    bold = next((run for run in caption_template.runs if run.bold), None)
    normal = next((run for run in caption_template.runs if run.text and not run.bold), None)
    if bold is None or normal is None:
        raise RuntimeError("Could not identify bold and normal caption runs in template")
    return bold, normal


def set_page_break_before(paragraph_element) -> None:
    properties = paragraph_element.get_or_add_pPr()
    for existing in properties.xpath("./w:pageBreakBefore"):
        properties.remove(existing)
    properties.append(OxmlElement("w:pageBreakBefore"))


def fixed_canvas(payload: bytes) -> bytes:
    with Image.open(io.BytesIO(payload)) as source:
        image = source.convert("RGBA")
        image.thumbnail(
            (FIGURE_FRAME_WIDTH_PX, FIGURE_FRAME_HEIGHT_PX),
            Image.Resampling.LANCZOS,
        )
        canvas = Image.new(
            "RGBA", (FIGURE_FRAME_WIDTH_PX, FIGURE_FRAME_HEIGHT_PX), "white"
        )
        x = (canvas.width - image.width) // 2
        y = (canvas.height - image.height) // 2
        canvas.alpha_composite(image, (x, y))
        output = io.BytesIO()
        canvas.convert("RGB").save(output, format="PNG", optimize=True)
        return output.getvalue()


def add_picture_paragraph(document, source, image_path: Path):
    return add_picture_bytes(
        document,
        source,
        fixed_canvas(image_path.read_bytes()),
        FIGURE_FRAME_WIDTH_IN,
        FIGURE_FRAME_HEIGHT_IN,
    )


def add_picture_bytes(document, source, payload: bytes, width: float, height: float):
    shape = document.add_picture(io.BytesIO(payload), width=Inches(width), height=Inches(height))
    element = shape._inline
    while not element.tag.endswith("}p"):
        element = element.getparent()
    paragraph = next(p for p in document.paragraphs if p._p is element)
    copy_paragraph_properties(source, paragraph)
    return paragraph


def add_caption(document, caption_template, heading_run, body_run, heading: str, body: str):
    paragraph = document.add_paragraph()
    copy_paragraph_properties(caption_template, paragraph)
    head = paragraph.add_run(heading)
    copy_run_properties(heading_run, head)
    if body:
        normal = paragraph.add_run(" " + body)
        copy_run_properties(body_run, normal)
    return paragraph


def find_template_blocks(document):
    paragraphs = document.paragraphs
    captions = {}
    images = {}
    for index, paragraph in enumerate(paragraphs):
        match = FIGURE_RE.match(paragraph.text.strip())
        if not match:
            continue
        number = int(match.group(1))
        captions[number] = copy.deepcopy(paragraph._p)
        for prior in reversed(paragraphs[:index]):
            if contains_drawing(prior._p):
                images[number] = copy.deepcopy(prior._p)
                break
        if number not in images:
            raise RuntimeError(f"No image paragraph found for original Figure S{number}")
    if sorted(captions) != list(range(1, 14)):
        raise RuntimeError(f"Expected original captions S1-S13, found {sorted(captions)}")
    return images, captions


def remove_original_figure_section(document) -> None:
    body = document._element.body
    first = next(
        index
        for index, element in enumerate(body)
        if element.tag.endswith("}p") and contains_drawing(element)
    )
    for element in list(body)[first:]:
        if not element.tag.endswith("}sectPr"):
            body.remove(element)


def add_cloned_block(
    document,
    image_template,
    image_payload: bytes,
    caption_element,
    old_number: int,
    new_number: int,
):
    image = add_picture_bytes(
        document,
        image_template,
        fixed_canvas(image_payload),
        FIGURE_FRAME_WIDTH_IN,
        FIGURE_FRAME_HEIGHT_IN,
    )
    set_page_break_before(image._p)
    caption = copy.deepcopy(caption_element)
    replace_prefix(caption, f"Figure S{old_number}.", f"Figure S{new_number}.")
    document._element.body.insert(-1, caption)


def chromosome_key(record: dict) -> tuple[int, int]:
    chrom = record["chrom"].removeprefix("chr")
    rank = {"X": 23, "Y": 24}.get(chrom, int(chrom) if chrom.isdigit() else 99)
    return rank, int(record["start"])


def consensus_loci() -> list[dict[str, str]]:
    payload = json.loads((REPO / "data/chimp_alignment_responses.json").read_text())
    records = sorted(payload["responses"], key=chromosome_key)
    properties = {}
    with (REPO / "data/inv_properties.tsv").open(newline="") as handle:
        for row in csv.DictReader(handle, delimiter="\t"):
            consensus = row["0_single_1_recur_consensus"].strip()
            if consensus in {"0", "1"}:
                properties[row["OrigID"]] = (
                    "Single-event" if consensus == "0" else "Recurrent"
                )
    selected = []
    for record in records:
        recurrence = properties.get(record["inv_id"])
        if recurrence:
            selected.append({**record, "recurrence": recurrence})
    if len(selected) != 93:
        raise RuntimeError(f"Expected 93 consensus loci, found {len(selected)}")
    return selected


def render_pdf_region_png(
    page, clip: pymupdf.Rect, target_width: int = 2400
) -> tuple[bytes, int, int]:
    scale = target_width / clip.width
    pixmap = page.get_pixmap(
        matrix=pymupdf.Matrix(scale, scale),
        clip=clip,
        alpha=False,
    )
    image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue(), image.width, image.height


def insert_front_matter_appendix_line(document, figure_range_paragraph, reference_paragraph) -> None:
    paragraph = document.add_paragraph()
    copy_paragraph_properties(reference_paragraph, paragraph)
    normal_source = next(run for run in reference_paragraph.runs if run.text)
    run = paragraph.add_run("SVbyEye alignments for 93 consensus-classified inversions")
    copy_run_properties(normal_source, run)
    figure_range_paragraph._p.addnext(paragraph._p)


def append_svbyeye(document, pdf_path: Path, image_template, caption_template, heading_run, body_run):
    source = pymupdf.open(pdf_path)
    loci = consensus_loci()
    expected_pages = (len(loci) + SVBYEYE_PLOTS_PER_PAGE - 1) // SVBYEYE_PLOTS_PER_PAGE
    if source.page_count != expected_pages:
        raise RuntimeError(
            f"SVbyEye PDF has {source.page_count} pages; expected {expected_pages} "
            "with two plots per page"
        )

    appendix_heading = add_caption(
        document,
        caption_template,
        heading_run,
        body_run,
        SVBYEYE_APPENDIX_TITLE,
        "Chimpanzee (panTro6) versus human GRCh38 alignments used for orientation polarization.",
    )
    set_page_break_before(appendix_heading._p)

    expected_page_size = None
    expected_pixels = None
    for index, locus in enumerate(loci, start=1):
        page = source[(index - 1) // SVBYEYE_PLOTS_PER_PAGE]
        page_size = (page.rect.width, page.rect.height)
        if expected_page_size is None:
            expected_page_size = page_size
            if page.rect.height <= page.rect.width / 2:
                raise RuntimeError(f"Unexpected paired SVbyEye page geometry: {page_size}")
        elif page_size != expected_page_size:
            raise RuntimeError(
                f"SVbyEye page size changed at alignment {index}: "
                f"expected {expected_page_size}, found {page_size}"
            )
        half_height = page.rect.height / SVBYEYE_PLOTS_PER_PAGE
        position = (index - 1) % SVBYEYE_PLOTS_PER_PAGE
        clip = pymupdf.Rect(
            0,
            position * half_height,
            page.rect.width,
            (position + 1) * half_height,
        )
        payload, width_px, height_px = render_pdf_region_png(page, clip)
        pixels = (width_px, height_px)
        if expected_pixels is None:
            expected_pixels = pixels
        elif pixels != expected_pixels:
            raise RuntimeError(
                f"SVbyEye page {index} rendered at {pixels}; expected {expected_pixels}"
            )
        image_paragraph = add_picture_bytes(
            document,
            image_template,
            payload,
            SVBYEYE_FRAME_WIDTH_IN,
            SVBYEYE_FRAME_WIDTH_IN * height_px / width_px,
        )
        if index > 1 and index % SVBYEYE_PLOTS_PER_PAGE == 1:
            set_page_break_before(image_paragraph._p)
        call = CALL_LABELS[locus["classification"]]
        heading = (
            f"SVbyEye alignment {index} of 93. "
            f"{locus['inv_id']} ({locus['region']})."
        )
        body = (
            "Chimpanzee (panTro6) versus human GRCh38 alignment used to polarize "
            f"orientation. Recurrence class: {locus['recurrence']}; GRCh38 "
            f"orientation relative to chimpanzee: {call}."
        )
        add_caption(document, caption_template, heading_run, body_run, heading, body)
    source.close()


def assemble(template: Path, output: Path, svbyeye_pdf: Path) -> None:
    if template.resolve() == output.resolve():
        raise ValueError("Output must differ from the immutable template")
    if sha256(template) != SUPPLEMENT_TEMPLATE_SHA256:
        raise RuntimeError(f"Template checksum does not match: {template}")
    source_checksum = sha256(template)

    document = Document(template)
    image_blocks, caption_blocks = find_template_blocks(document)
    image_template = next(
        paragraph for paragraph in document.paragraphs if contains_drawing(paragraph._p)
    )
    caption_template = next(
        paragraph for paragraph in document.paragraphs if FIGURE_RE.match(paragraph.text.strip())
    )
    heading_run, body_run = template_runs(caption_template)
    image_payloads = {}
    for number, element in image_blocks.items():
        blips = element.xpath(".//a:blip")
        if len(blips) != 1:
            raise RuntimeError(f"Expected one embedded image for original Figure S{number}")
        relationship = blips[0].get(qn("r:embed"))
        image_payloads[number] = document.part.related_parts[relationship].blob

    figure_range = next(p for p in document.paragraphs if p.text.strip() == "Figs. S1 to S13")
    table_range = next(p for p in document.paragraphs if p.text.strip() == "Tables S1 to S20")
    replace_paragraph_text(figure_range, "Figs. S1 to S13", "Figs. S1 to S22")
    replace_paragraph_text(table_range, "Tables S1 to S20", "Tables S1 to S21")
    insert_front_matter_appendix_line(document, figure_range, figure_range)

    remove_original_figure_section(document)

    for figure in FINAL_SUPPLEMENTARY_FIGURES:
        if figure.original_number is not None:
            add_cloned_block(
                document,
                image_template,
                image_payloads[figure.original_number],
                caption_blocks[figure.original_number],
                figure.original_number,
                figure.number,
            )
            continue
        if not figure.asset or figure.caption is None:
            raise RuntimeError(f"Incomplete revision-figure metadata: S{figure.number}")
        asset = REPO / figure.asset
        if not asset.is_file():
            raise FileNotFoundError(asset)
        image_paragraph = add_picture_paragraph(document, image_template, asset)
        set_page_break_before(image_paragraph._p)
        add_caption(
            document,
            caption_template,
            heading_run,
            body_run,
            f"Figure S{figure.number}. {figure.title}.",
            figure.caption,
        )

    append_svbyeye(
        document,
        svbyeye_pdf,
        image_template,
        caption_template,
        heading_run,
        body_run,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        dir=output.parent, prefix=f".{output.name}.", suffix=".docx", delete=False
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        restore_template_parts(temporary, template)
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)

    if sha256(template) != source_checksum:
        raise RuntimeError("Immutable template changed during assembly")
    print(f"Assembled {output}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=REPO / SUPPLEMENT_TEMPLATE)
    parser.add_argument("--svbyeye-pdf", type=Path, default=REPO / SVBYEYE_CONSENSUS_PDF)
    parser.add_argument(
        "--output",
        type=Path,
        default=REPO / "output/supplementary_materials/Supplementary_Materials_final.docx",
    )
    args = parser.parse_args()
    assemble(args.template, args.output, args.svbyeye_pdf)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
