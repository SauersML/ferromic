#!/usr/bin/env python3
"""Verify the final supplementary-material DOCX structure and ordering."""

from __future__ import annotations

import argparse
import hashlib
import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document


REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "stats"))

from supplementary_inventory import (  # noqa: E402
    FINAL_SUPPLEMENTARY_FIGURES,
    RESPONSE_ONLY_FIGURE_TITLES,
    SUPPLEMENT_TEMPLATE,
    SUPPLEMENT_TEMPLATE_SHA256,
    SVBYEYE_APPENDIX_TITLE,
)


FIGURE_RE = re.compile(r"^Figure S(\d+)\.")
APPENDIX_RE = re.compile(r"^SVbyEye alignment (\d+) of 93\. (\S+) \(")
IMMUTABLE_PARTS = (
    "word/styles.xml",
    "word/settings.xml",
    "word/fontTable.xml",
    "word/theme/theme1.xml",
    "word/header1.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def verify(template: Path, assembled: Path) -> None:
    if template.resolve() == assembled.resolve():
        raise RuntimeError("Assembled output points to the immutable template")
    if sha256(template) != SUPPLEMENT_TEMPLATE_SHA256:
        raise RuntimeError("Immutable template checksum mismatch")

    with ZipFile(assembled) as archive:
        corrupt = archive.testzip()
        if corrupt:
            raise RuntimeError(f"Corrupt DOCX member: {corrupt}")
        with ZipFile(template) as source:
            for member in IMMUTABLE_PARTS:
                if source.read(member) != archive.read(member):
                    raise RuntimeError(f"Template formatting part changed: {member}")
            for name in (f"word/media/image{i}.png" for i in range(1, 14)):
                if source.read(name) != archive.read(name):
                    raise RuntimeError(f"Original figure image bytes changed: {name}")

    document = Document(assembled)
    texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    full_text = "\n".join(texts)

    captions = [text for text in texts if FIGURE_RE.match(text)]
    numbers = [int(FIGURE_RE.match(text).group(1)) for text in captions]
    if numbers != list(range(1, 23)):
        raise RuntimeError(f"Expected exactly ordered captions S1-S22; found {numbers}")
    for caption, figure in zip(captions, FINAL_SUPPLEMENTARY_FIGURES):
        expected = f"Figure S{figure.number}. {figure.title}"
        if not caption.startswith(expected):
            raise RuntimeError(
                f"Figure S{figure.number} title mismatch:\n"
                f"expected prefix={expected!r}\nobserved={caption[:180]!r}"
            )

    promoted = [title for title in RESPONSE_ONLY_FIGURE_TITLES if title in full_text]
    if promoted:
        raise RuntimeError(f"Response-only figures were promoted into the supplement: {promoted}")
    if "Figs. S1 to S22" not in texts or "Tables S1 to S21" not in texts:
        raise RuntimeError("Front-matter figure/table ranges were not updated")
    if "SVbyEye alignments for 93 consensus-classified inversions" not in texts:
        raise RuntimeError("Front matter does not list the SVbyEye appendix")
    if not any(text.startswith(SVBYEYE_APPENDIX_TITLE) for text in texts):
        raise RuntimeError("SVbyEye appendix title is absent")

    appendix = [APPENDIX_RE.match(text) for text in texts if APPENDIX_RE.match(text)]
    indices = [int(match.group(1)) for match in appendix]
    if indices != list(range(1, 94)):
        raise RuntimeError(f"Expected 93 ordered SVbyEye captions; found {indices}")
    if len({match.group(2) for match in appendix}) != 93:
        raise RuntimeError("SVbyEye appendix locus identifiers are not unique")

    drawings = document._element.body.xpath(".//w:drawing")
    if len(drawings) != 115:
        raise RuntimeError(f"Expected 22 + 93 = 115 drawings; found {len(drawings)}")
    figure_page_starts = document._element.body.xpath(
        ".//w:p[w:pPr/w:pageBreakBefore][.//w:drawing]"
    )
    if len(figure_page_starts) != 68:
        raise RuntimeError(
            "Expected page-break-before on 22 numbered figures and 46 subsequent "
            f"appendix plots; found {len(figure_page_starts)}"
        )
    standalone_page_breaks = document._element.body.xpath(
        ".//w:br[@w:type='page']"
    )
    if standalone_page_breaks:
        raise RuntimeError(
            "Standalone page-break paragraphs can create blank rendered pages; "
            f"found {len(standalone_page_breaks)}"
        )
    if len(document.sections) != 1:
        raise RuntimeError(
            "The supplement must retain one original-format portrait section; "
            f"found {len(document.sections)} sections"
        )
    template_document = Document(template)
    expected_section = template_document.sections[0]
    observed_section = document.sections[0]
    page_geometry = (
        "page_width",
        "page_height",
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
        "header_distance",
        "footer_distance",
    )
    for attribute in page_geometry:
        expected = getattr(expected_section, attribute)
        observed = getattr(observed_section, attribute)
        if observed != expected:
            raise RuntimeError(
                f"Original page geometry changed for {attribute}: "
                f"expected {expected}, observed {observed}"
            )

    extents = document._element.body.xpath(".//w:drawing/wp:inline/wp:extent")
    figure_extents = [
        (int(extent.get("cx")), int(extent.get("cy"))) for extent in extents
    ]
    main_figure_extents = figure_extents[:22]
    appendix_extents = figure_extents[22:]
    if (
        len(figure_extents) != 115
        or len(set(main_figure_extents)) != 1
        or len(set(appendix_extents)) != 1
        or main_figure_extents[0][0] != appendix_extents[0][0]
        or main_figure_extents[0][1] <= appendix_extents[0][1]
    ):
        raise RuntimeError(
            "Expected one fixed frame for the 22 numbered figures and one shorter, "
            "full-width frame for the 93 SVbyEye plots; found "
            f"main={sorted(set(main_figure_extents))}, "
            f"appendix={sorted(set(appendix_extents))}"
        )

    print(
        "Verified assembled supplement: S1-S22 ordered; response-only figures omitted; "
        "all 93 SVbyEye drawings use a full-width two-per-page frame; 93-caption "
        "portrait appendix present; template formatting parts and page geometry unchanged."
    )


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--template", type=Path, default=REPO / SUPPLEMENT_TEMPLATE)
    args = parser.parse_args()
    verify(args.template, args.docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
